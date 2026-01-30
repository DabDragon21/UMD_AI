import os
import streamlit as st
from io import BytesIO
from docx import Document  # pip install python-docx

# -----------------------
# ✅ Set API key
# -----------------------
OPENAI_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_KEY:
    try:
        OPENAI_KEY = st.secrets["OPENAI_KEY"]
    except Exception:
        import getpass

        OPENAI_KEY = getpass.getpass("Enter your OpenAI API key:")

os.environ["OPENAI_API_KEY"] = OPENAI_KEY

# -----------------------
# ✅ Imports (updated for latest LangChain)
# -----------------------
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyMuPDFLoader
from langchain.messages import HumanMessage

# -----------------------
# App start
# -----------------------
st.set_page_config(page_title="AI Lesson Assistant", layout="centered")
st.title("📚 AI Lesson Assistant")
st.write("LangChain + OpenAI embeddings loaded")

# -----------------------
# Load PDFs and vectorstore
# -----------------------
databases = [
    "Differentiation_Guides.pdf",
    "MaCann Yadav CT elem.pdf",
    "MultilingualLearnersGuide.pdf",
    "UDL_Table_accessible_CS.pdf",
    "CS_Pedagogy.pdf",
    "CS_Content.pdf"
]

vectorstore_path = "edu_faiss_index"
if not os.path.exists(vectorstore_path):
    os.makedirs(vectorstore_path)

embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
vectorstore = None

if os.path.exists(vectorstore_path):
    st.write("🔄 Loading cached vectorstore...")
    vectorstore = FAISS.load_local(vectorstore_path, embeddings, allow_dangerous_deserialization=True)
else:
    st.write("⚙️ Building new vectorstore...")
    all_docs = []
    for file in databases:
        loader = PyMuPDFLoader(file)
        docs = loader.load()
        all_docs.extend(docs)

    splitter = RecursiveCharacterTextSplitter(chunk_size=256, chunk_overlap=64)
    chunks = splitter.split_documents(all_docs)
    chunks = chunks[:50]  # optional limit

    vectorstore = FAISS.from_documents(chunks, embeddings)
    vectorstore.save_local(vectorstore_path)
    st.write("Vectorstore saved ✅")

# -----------------------
# Setup LLM
# -----------------------
LLM = ChatOpenAI(model_name="gpt-3.5-turbo", temperature=0)


# -----------------------
# Build simple RAG function
# -----------------------
def run_rag(vectorstore, prompt_text):
    docs = vectorstore.similarity_search(prompt_text, k=4)
    context_text = "\n\n".join([doc.page_content for doc in docs])

    final_prompt = f"""
You are an expert educational assistant.
Use the following research excerpts to improve the lesson plan.
If the research is irrelevant, say so explicitly.

Research context:
{context_text}

Lesson plan and request:
{prompt_text}
"""

    messages = [HumanMessage(content=final_prompt)]
    response = LLM.generate([messages])
    return response.generations[0][0].text


# -----------------------
# Streamlit UI
# -----------------------
st.markdown("Upload your educational PDF and enter a prompt to update the lesson plan.")

uploaded_file = st.file_uploader("Choose a PDF file", type=["pdf"])
prompt = st.text_area("Enter your prompt", height=100)

if uploaded_file and prompt:
    if st.button("Generate Updated Lesson Plan"):
        with st.spinner("Processing..."):
            # Save uploaded PDF temporarily
            with open("temp_uploaded.pdf", "wb") as f:
                f.write(uploaded_file.read())

            loader = PyMuPDFLoader("temp_uploaded.pdf")
            docs = loader.load()
            input_txt = "".join([doc.page_content for doc in docs])
            truncated = input_txt[:8000]

            # Run RAG
            output_text = run_rag(vectorstore, f"{prompt}\n\nLesson plan:\n{truncated}")

            # Show in text area
            st.text_area("Updated Lesson Plan", output_text, height=300)

            # -----------------------
            # Create downloadable DOCX
            # -----------------------
            doc = Document()
            for line in output_text.split("\n"):
                line = line.strip()
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line)

            # Save to in-memory bytes
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Streamlit download button
            st.download_button(
                label="📄 Download Updated Lesson Plan (.docx)",
                data=buffer,
                file_name="updated_lesson_plan.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
